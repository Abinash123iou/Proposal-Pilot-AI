import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt
from app.core.logger import logger
from app.core.config import settings, BASE_DIR
from app.document.generator import DocumentGenerator
from app.document.styles import (
    COLOR_PRIMARY,
    COLOR_SECONDARY,
    COLOR_WARNING,
    COLOR_TEXT,
    PRIMARY_FONT,
    SECONDARY_FONT,
    SIZE_TITLE,
    SIZE_SUBTITLE,
    SIZE_H1,
    MARGIN_INCHES
)
from app.document.formatter import (
    apply_run_font,
    apply_paragraph_spacing,
    set_document_margins,
    setup_header_footer,
    add_markdown_paragraph
)
from app.document.templates import STANDARD_SECTIONS, SECTION_SYNONYMS

class DocxDocumentGenerator(DocumentGenerator):
    """
    Concrete implementation of DocumentGenerator that produces Microsoft Word (.docx) documents.
    """
    
    async def generate(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """
        Builds a professionally styled DOCX file from the validated AgentState and saves it.
        """
        logger.info("Document generation started")
        errors: List[str] = []
        
        # 1. Input Validation
        if not isinstance(state, dict):
            err = "Invalid AgentState: state must be a dictionary."
            logger.error(err)
            return {"success": False, "errors": [err]}
            
        generated_sections = state.get("generated_sections")
        if not generated_sections:
            err = "Document generation failed: generated_sections is empty or missing."
            logger.error(err)
            return {"success": False, "errors": [err]}
            
        # 2. Output Directory Configuration Check
        try:
            output_dir = Path(BASE_DIR) / settings.OUTPUT_DIR
            os.makedirs(output_dir, exist_ok=True)
        except Exception as e:
            err = f"Failed to configure output directory: {str(e)}"
            logger.error(err)
            return {"success": False, "errors": [err]}
            
        # 3. File Naming & Overwrite Prevention
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"Proposal_{timestamp}"
        filename = f"{base_filename}.docx"
        file_path = output_dir / filename
        
        counter = 1
        while file_path.exists():
            filename = f"{base_filename}_{counter}.docx"
            file_path = output_dir / filename
            counter += 1
            
        try:
            # 4. Initialize Document structure
            doc = Document()
            set_document_margins(doc, MARGIN_INCHES)
            
            # --- Cover Page Layout ---
            logger.info("Writing Cover Page")
            p_spacer = doc.add_paragraph()
            apply_paragraph_spacing(p_spacer, space_before=Pt(72))
            
            # Title
            p_title = doc.add_paragraph()
            apply_paragraph_spacing(p_title, space_after=Pt(12))
            run_title = p_title.add_run("PROJECT PROPOSAL")
            apply_run_font(run_title, PRIMARY_FONT, SIZE_TITLE, COLOR_PRIMARY, bold=True)
            
            # Subtitle
            p_sub = doc.add_paragraph()
            apply_paragraph_spacing(p_sub, space_after=Pt(36))
            request_text = state.get("request", "Consulting & Services Agreement")
            run_sub = p_sub.add_run(request_text)
            apply_run_font(run_sub, SECONDARY_FONT, SIZE_SUBTITLE, COLOR_SECONDARY, italic=True)
            
            # Divider line
            p_div = doc.add_paragraph()
            apply_paragraph_spacing(p_div, space_after=Pt(120))
            run_div = p_div.add_run("—" * 40)
            apply_run_font(run_div, PRIMARY_FONT, Pt(12), COLOR_SECONDARY)
            
            # Meta details
            p_meta = doc.add_paragraph()
            apply_paragraph_spacing(p_meta, space_after=Pt(6))
            
            run_org = p_meta.add_run("Prepared for:\nClient Organization\n\n")
            apply_run_font(run_org, PRIMARY_FONT, Pt(11), COLOR_SECONDARY)
            
            run_author = p_meta.add_run("Prepared by: ProposalPilot AI\n")
            apply_run_font(run_author, PRIMARY_FONT, Pt(11), COLOR_SECONDARY)
            
            quality_score = state.get("quality_score", 0)
            if quality_score:
                run_score = p_meta.add_run(f"Quality QA Score: {quality_score}/100\n")
                apply_run_font(run_score, PRIMARY_FONT, Pt(10), COLOR_SECONDARY, bold=True)
                
            run_date = p_meta.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            apply_run_font(run_date, PRIMARY_FONT, Pt(10), COLOR_SECONDARY)
            
            doc.add_page_break()
            
            # --- Table of Contents Placeholder ---
            logger.info("Writing Table of Contents Page")
            p_toc_head = doc.add_paragraph(style='Heading 1')
            apply_paragraph_spacing(p_toc_head, space_before=Pt(12), space_after=Pt(18))
            run_toc_head = p_toc_head.add_run("Table of Contents")
            apply_run_font(run_toc_head, PRIMARY_FONT, SIZE_H1, COLOR_PRIMARY, bold=True)
            
            p_toc_desc = doc.add_paragraph()
            apply_paragraph_spacing(p_toc_desc, space_after=Pt(12))
            run_toc_desc = p_toc_desc.add_run(
                "This document is structured into the following sections. "
                "The page numbers will align dynamically upon opening in Microsoft Word."
            )
            apply_run_font(run_toc_desc, PRIMARY_FONT, Pt(11), COLOR_TEXT, italic=True)
            
            for sec_name in STANDARD_SECTIONS:
                p_toc_item = doc.add_paragraph(style='List Bullet')
                apply_paragraph_spacing(p_toc_item, space_after=Pt(4))
                run_item = p_toc_item.add_run(sec_name)
                apply_run_font(run_item, PRIMARY_FONT, Pt(11), COLOR_TEXT)
                
            doc.add_page_break()
            
            # --- Content Sectioning Loop ---
            logger.info("Sections written")
            for sec_name in STANDARD_SECTIONS:
                p_h1 = doc.add_paragraph(style='Heading 1')
                apply_paragraph_spacing(p_h1, space_before=Pt(24), space_after=Pt(12))
                run_h1 = p_h1.add_run(sec_name)
                apply_run_font(run_h1, PRIMARY_FONT, SIZE_H1, COLOR_PRIMARY, bold=True)
                
                # Locate text fuzzy matching synonyms
                content = None
                synonyms = SECTION_SYNONYMS.get(sec_name, [sec_name.lower()])
                for key in generated_sections:
                    key_lower = key.lower()
                    if any(synonym in key_lower for synonym in synonyms):
                        content = generated_sections[key]
                        break
                        
                if content and content.strip():
                    paragraphs = content.split("\n")
                    for p_text in paragraphs:
                        if p_text.strip():
                            add_markdown_paragraph(doc, p_text)
                else:
                    # Missing section handling
                    logger.warning(f"Section '{sec_name}' is missing in state's generated sections.")
                    errors.append(f"Missing section: '{sec_name}'")
                    p_warn = doc.add_paragraph()
                    apply_paragraph_spacing(p_warn, space_after=Pt(6))
                    run_warn = p_warn.add_run(f"[Content for section '{sec_name}' is missing.]")
                    apply_run_font(run_warn, PRIMARY_FONT, Pt(11), COLOR_WARNING, italic=True, bold=True)
            
            logger.info("Formatting applied")
            # Apply Header & Footer styling
            setup_header_footer(doc, "Project Proposal  |  Confidential")
            
            # 5. Save output
            logger.info("File saved")
            doc.save(file_path)
            
            size_bytes = os.path.getsize(file_path)
            logger.info(f"File size: {size_bytes}")
            logger.info("Document generation completed")
            
            return {
                "success": True,
                "document_path": str(file_path),
                "filename": filename,
                "file_size_bytes": size_bytes,
                "timestamp": timestamp,
                "errors": errors
            }
            
        except PermissionError as e:
            err_msg = f"File permission error while saving document: {str(e)}"
            logger.error(err_msg)
            return {"success": False, "errors": [err_msg]}
            
        except Exception as e:
            err_msg = f"DOCX generation failures: {str(e)}"
            logger.error(err_msg)
            return {"success": False, "errors": [err_msg]}

import re
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.document.styles import (
    PRIMARY_FONT,
    COLOR_TEXT,
    SPACING_LINE,
    SPACING_PARAGRAPH_AFTER,
    SIZE_BODY
)

def apply_run_font(run, font_name: str, font_size: Pt, color: RGBColor, bold: bool = False, italic: bool = False):
    """
    Applies font details directly to a text run.
    """
    run.font.name = font_name
    run.font.size = font_size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic

def apply_paragraph_spacing(paragraph, space_before: Pt = Pt(0), space_after: Pt = Pt(6), line_spacing: float = 1.15):
    """
    Applies spacing details to a paragraph.
    """
    p_format = paragraph.paragraph_format
    p_format.space_before = space_before
    p_format.space_after = space_after
    p_format.line_spacing = line_spacing

def set_document_margins(doc, margin_value):
    """
    Configures standard margins for all sections of the document.
    """
    for section in doc.sections:
        section.top_margin = margin_value
        section.bottom_margin = margin_value
        section.left_margin = margin_value
        section.right_margin = margin_value

def add_page_number_to_run(run):
    """
    Appends an XML page number field into a run.
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def setup_header_footer(doc, header_text: str):
    """
    Applies right-aligned running headers and centered page numbering in footers.
    Uses 'different_first_page_header_footer' to suppress these on the cover page.
    """
    for section in doc.sections:
        section.different_first_page_header_footer = True
        
        # Configure Header (pages 2+)
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = header_text
        header_p.alignment = 2  # Right-aligned
        apply_paragraph_spacing(header_p, space_before=Pt(0), space_after=Pt(4))
        for run in header_p.runs:
            apply_run_font(run, PRIMARY_FONT, Pt(9), RGBColor(128, 128, 128), italic=True)
            
        # Configure Footer (pages 2+)
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = "ProposalPilot AI  |  "
        apply_paragraph_spacing(footer_p, space_before=Pt(4), space_after=Pt(0))
        footer_p.alignment = 1  # Centered
        
        for run in footer_p.runs:
            apply_run_font(run, PRIMARY_FONT, Pt(9), RGBColor(128, 128, 128), italic=True)
            
        page_run = footer_p.add_run("Page ")
        apply_run_font(page_run, PRIMARY_FONT, Pt(9), RGBColor(128, 128, 128), italic=True)
        num_run = footer_p.add_run()
        apply_run_font(num_run, PRIMARY_FONT, Pt(9), RGBColor(128, 128, 128), italic=True)
        add_page_number_to_run(num_run)

def add_markdown_paragraph(doc, text: str):
    """
    Parses bullet points (- / *), numbered lists (1. ), bold (**), and italic (*) 
    markdown notations from input text and adds them as stylized paragraphs and runs.
    """
    text_stripped = text.strip()
    if not text_stripped:
        p = doc.add_paragraph()
        apply_paragraph_spacing(p, space_after=SPACING_PARAGRAPH_AFTER, line_spacing=SPACING_LINE)
        return p
        
    style_name = 'Normal'
    
    # Identify lists
    if text_stripped.startswith("- ") or text_stripped.startswith("* "):
        style_name = 'List Bullet'
        text_stripped = text_stripped[2:]
    elif re.match(r'^\d+\.\s', text_stripped):
        style_name = 'List Number'
        match = re.match(r'^(\d+\.\s)', text_stripped)
        prefix_len = len(match.group(1))
        text_stripped = text_stripped[prefix_len:]
        
    p = doc.add_paragraph(style=style_name)
    apply_paragraph_spacing(p, space_after=SPACING_PARAGRAPH_AFTER, line_spacing=SPACING_LINE)
    
    # Parsing bold and italics
    bold_parts = re.split(r'(\*\*.*?\*\*)', text_stripped)
    for b_part in bold_parts:
        is_bold = b_part.startswith('**') and b_part.endswith('**')
        clean_b_part = b_part[2:-2] if is_bold else b_part
        
        italic_parts = re.split(r'(\*.*?\*)', clean_b_part)
        for i_part in italic_parts:
            is_italic = i_part.startswith('*') and i_part.endswith('*')
            clean_i_part = i_part[1:-1] if is_italic else i_part
            
            if not clean_i_part:
                continue
                
            run = p.add_run(clean_i_part)
            apply_run_font(
                run=run,
                font_name=PRIMARY_FONT,
                font_size=SIZE_BODY,
                color=COLOR_TEXT,
                bold=is_bold,
                italic=is_italic
            )
            
    return p

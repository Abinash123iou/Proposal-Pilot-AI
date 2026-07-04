import os
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock
from datetime import datetime
from docx import Document
from app.core.config import settings, BASE_DIR
from app.document.docx_generator import DocxDocumentGenerator

# Valid sample content for all 12 proposal sections
SAMPLE_SECTIONS = {
    "Executive Summary": "This is a brief summary of the proposal, detailing standard guidelines.",
    "Client Requirements": "The client requires an enterprise-scale AI assistant.",
    "Project Objectives": "Deliver system on time, under budget, and with high QA scores.",
    "Scope of Work": "Implement frontend, backend, database models, and workflow tests.",
    "Functional Modules": "Planner, Executor, Reflection, and Document Generation.",
    "Technology Stack": "FastAPI, LangGraph, python-docx, groq, pytest.",
    "System Architecture": "Autonomous agent pattern with stateful LangGraph workflow.",
    "Project Timeline": "Weeks 1-4 planning; Weeks 5-8 execution; Week 9 rollout.",
    "Budget Estimation": "Total estimation is $50,000 for standard delivery.",
    "Risks": "API rate limits, server downtime, LLM hallucinations.",
    "Assumptions": "Groq API is accessible and python environment is functional.",
    "Conclusion": "The proposal satisfies all client needs. Ready to proceed."
}

@pytest.fixture
def sample_state():
    return {
        "request": "Hospital Scheduling Software Proposal",
        "generated_sections": SAMPLE_SECTIONS.copy(),
        "quality_score": 95,
        "metadata": {"author": "Pilot"}
    }

@pytest.mark.anyio
async def test_generate_success(sample_state):
    """
    1. Successful document generation.
    Verifies margins, section headings, metadata, formatting, and file storage.
    """
    generator = DocxDocumentGenerator()
    result = await generator.generate(sample_state)
    
    assert result["success"] is True
    assert "document_path" in result
    assert result["filename"].startswith("Proposal_")
    assert result["file_size_bytes"] > 0
    assert len(result["errors"]) == 0
    
    file_path = Path(result["document_path"])
    assert file_path.exists()
    
    # Verify contents using python-docx
    doc = Document(file_path)
    
    # Verify margins (1 inch)
    for section in doc.sections:
        assert abs(section.top_margin.inches - 1.0) < 1e-4
        assert abs(section.bottom_margin.inches - 1.0) < 1e-4
        
    # Verify heading hierarchy and paragraphs
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Table of Contents" in headings
    for sec_name in SAMPLE_SECTIONS.keys():
        assert sec_name in headings
        
    # Clean up
    if file_path.exists():
        os.remove(file_path)

@pytest.mark.anyio
async def test_generate_missing_sections(sample_state):
    """
    2. Missing section handling.
    Verifies that missing sections produce warning placeholders in the document
    and are captured in the return metadata errors list without crashing.
    """
    # Remove some sections
    del sample_state["generated_sections"]["Scope of Work"]
    del sample_state["generated_sections"]["Budget Estimation"]
    
    generator = DocxDocumentGenerator()
    result = await generator.generate(sample_state)
    
    assert result["success"] is True
    assert len(result["errors"]) == 2
    assert "Missing section: 'Scope of Work'" in result["errors"]
    assert "Missing section: 'Budget Estimation'" in result["errors"]
    
    file_path = Path(result["document_path"])
    doc = Document(file_path)
    
    # Verify placeholder warning text exists
    body_texts = [p.text for p in doc.paragraphs]
    assert "[Content for section 'Scope of Work' is missing.]" in body_texts
    assert "[Content for section 'Budget Estimation' is missing.]" in body_texts
    
    # Clean up
    if file_path.exists():
        os.remove(file_path)

@pytest.mark.anyio
async def test_generate_invalid_output_directory(sample_state):
    """
    3. Invalid output directory handling.
    Verifies that a directory creation failure or write exception returns a structured error.
    """
    generator = DocxDocumentGenerator()
    
    # Mock settings.OUTPUT_DIR to an invalid path containing invalid characters (\x00)
    with patch.object(settings, "OUTPUT_DIR", "invalid_dir\x00_path"):
        result = await generator.generate(sample_state)
        
        assert result["success"] is False
        assert len(result["errors"]) > 0
        assert "Failed to configure output directory" in result["errors"][0] or "DOCX generation" in result["errors"][0]

@pytest.mark.anyio
async def test_generate_empty_proposal(sample_state):
    """
    4. Empty proposal handling.
    Verifies that passing an empty proposal or missing generated_sections field is rejected.
    """
    generator = DocxDocumentGenerator()
    
    # Empty generated_sections
    sample_state["generated_sections"] = {}
    result = await generator.generate(sample_state)
    assert result["success"] is False
    assert "generated_sections is empty" in result["errors"][0]
    
    # Non-dictionary state
    result = await generator.generate("not a dictionary")
    assert result["success"] is False
    assert "state must be a dictionary" in result["errors"][0]

@pytest.mark.anyio
async def test_generate_overwrite_prevention(sample_state):
    """
    5. File overwrite prevention.
    Verifies that generating multiple documents with the same timestamp does not overwrite,
    but appends a serial counter to preserve both files.
    """
    generator = DocxDocumentGenerator()
    
    fixed_now = datetime(2026, 7, 4, 15, 30, 0)
    
    # Mock datetime.now to return fixed timestamp for both invocations
    with patch("app.document.docx_generator.datetime") as mock_datetime:
        mock_datetime.now.return_value = fixed_now
        mock_datetime.strftime = datetime.strftime
        
        res1 = await generator.generate(sample_state)
        res2 = await generator.generate(sample_state)
        
        assert res1["success"] is True
        assert res2["success"] is True
        
        path1 = Path(res1["document_path"])
        path2 = Path(res2["document_path"])
        
        assert path1.exists()
        assert path2.exists()
        assert path1 != path2
        assert "Proposal_20260704_153000.docx" in res1["filename"]
        assert "Proposal_20260704_153000_1.docx" in res2["filename"]
        
        # Clean up
        if path1.exists():
            os.remove(path1)
        if path2.exists():
            os.remove(path2)
